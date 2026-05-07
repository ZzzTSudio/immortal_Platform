"""Ephemeral chat attachment parsing for prompt context."""

from __future__ import annotations

import asyncio
from io import BytesIO
from pathlib import Path
from typing import Callable

from fastapi import UploadFile

from app.rag.chunker import estimate_tokens
from app.rag.cleaner import clean_text

MAX_CHAT_FILES = 3
MAX_CHAT_FILE_BYTES = 5 * 1024 * 1024
MAX_FILE_CONTEXT_TOKENS = 2500
SUPPORTED_CHAT_FILE_SUFFIXES = {".pdf", ".md", ".doc", ".docx", ".txt"}

_READ_CHUNK_SIZE = 1024 * 1024

StatusCallback = Callable[[str], None]


class ChatFileValidationError(ValueError):
    """Raised when a chat upload violates hard limits."""


async def build_chat_file_context(
    files: list[UploadFile] | None,
    *,
    max_tokens: int = MAX_FILE_CONTEXT_TOKENS,
    on_status: StatusCallback | None = None,
) -> str:
    valid_files = _validate_files(files or [])
    blocks: list[tuple[str, str]] = []

    for upload in valid_files:
        filename = Path(upload.filename or "unnamed").name
        if on_status is not None:
            on_status(f"解析文件 {filename} 中…")
        data = await _read_limited(upload)
        text = await asyncio.to_thread(_extract_text_safely, filename, data)
        cleaned = clean_text(text)
        if cleaned.strip():
            blocks.append((filename, cleaned.strip()))

    return _format_file_context(blocks, max_tokens=max_tokens)


def _validate_files(files: list[UploadFile]) -> list[UploadFile]:
    if len(files) > MAX_CHAT_FILES:
        raise ChatFileValidationError(f"最多上传 {MAX_CHAT_FILES} 个文件")

    for upload in files:
        filename = Path(upload.filename or "").name
        suffix = Path(filename).suffix.lower()
        if suffix not in SUPPORTED_CHAT_FILE_SUFFIXES:
            raise ChatFileValidationError(f"不支持的文件格式：{suffix or filename}")
    return files


async def _read_limited(upload: UploadFile) -> bytes:
    chunks: list[bytes] = []
    total = 0
    while True:
        chunk = await upload.read(_READ_CHUNK_SIZE)
        if not chunk:
            break
        total += len(chunk)
        if total > MAX_CHAT_FILE_BYTES:
            raise ChatFileValidationError("单个文件大小不能超过 5MB")
        chunks.append(chunk)
    return b"".join(chunks)


def _extract_text_safely(filename: str, data: bytes) -> str:
    try:
        suffix = Path(filename).suffix.lower()
        if suffix == ".pdf":
            return _extract_pdf(data)
        if suffix == ".docx":
            return _extract_docx(data)
        if suffix in {".md", ".txt"}:
            return _decode_text(data)
        if suffix == ".doc":
            return _extract_doc_best_effort(data)
    except Exception:
        return ""
    return ""


def _decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "gbk"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(data))
    pages = [(page.extract_text() or "") for page in reader.pages]
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    from docx import Document

    document = Document(BytesIO(data))
    parts: list[str] = []
    parts.extend(p.text for p in document.paragraphs if p.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_doc_best_effort(data: bytes) -> str:
    # Legacy .doc is a binary format. Without a system converter, only keep obvious text runs.
    text = data.decode("latin-1", errors="ignore")
    printable = "".join(ch if ch.isprintable() or ch in "\r\n\t" else " " for ch in text)
    lines = [line.strip() for line in printable.splitlines() if len(line.strip()) >= 8]
    return "\n".join(lines)


def _format_file_context(blocks: list[tuple[str, str]], *, max_tokens: int) -> str:
    if not blocks or max_tokens <= 0:
        return ""

    lines: list[str] = []
    total_tokens = 0
    for index, (filename, content) in enumerate(blocks, start=1):
        header = f"[FILE {index}: {filename}]"
        header_tokens = estimate_tokens(header)
        if total_tokens + header_tokens >= max_tokens:
            break

        remaining = max_tokens - total_tokens - header_tokens
        body, truncated = _truncate_to_tokens(content, remaining)
        if not body.strip():
            break

        lines.extend([header, body.strip()])
        total_tokens += header_tokens + estimate_tokens(body)
        if truncated:
            break
    return "\n".join(lines).strip()


def _truncate_to_tokens(text: str, max_tokens: int) -> tuple[str, bool]:
    if max_tokens <= 0:
        return "", True
    if estimate_tokens(text) <= max_tokens:
        return text, False
    max_chars = max(20, max_tokens * 2)
    return text[: max_chars - 8].rstrip() + "…(截断)", True
